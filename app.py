import streamlit as st
import io
import re
import os
import uuid
import tempfile
from docx import Document as DocxDocument

# ── Backend imports ──────────────────────────────────────────────────────────
from src.rag_pipeline import RAGPipeline
from src.document_processor import DocumentProcessor
from src.web_scraper import WebScraper
from src.web_search import WebSearchClient
from src.llm_router import LLMRouter
from src.conversation_memory import ConversationMemory
from streamlit_mic_recorder import mic_recorder
from src.voice_assistant import VoiceAssistant

# ==============================================================================
# 1. PAGE CONFIGURATION
# ==============================================================================
st.set_page_config(
    page_title="Alies AI ChatBot RAG System",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .stChatInputContainer { padding-bottom: 20px; }
    div[data-testid="stSidebarUserContent"] { padding-top: 1rem; }
    div[data-testid="stPopover"] button svg { display: none; }
    .source-tag {
        background-color: #fff3e0;
        color: #f57c00;
        padding: 2px 8px;
        border-radius: 4px;
        font-size: 0.78rem;
        margin: 2px 2px;
        display: inline-block;
    }
    .code-label {
        background-color: #1e1e1e;
        color: #9cdcfe;
        padding: 3px 10px;
        border-radius: 5px 5px 0 0;
        font-size: 0.73rem;
        font-family: monospace;
        display: inline-block;
        margin-bottom: -4px;
    }
    .backend-badge { font-size: 0.72rem; color: #888; margin-top: 4px; }
    .mode-badge {
        padding: 4px 10px;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: 600;
        display: inline-block;
        margin-bottom: 8px;
    }
</style>
""", unsafe_allow_html=True)

# ==============================================================================
# 2. SESSION STATE INITIALIZATION
# ==============================================================================
def init_state():
    if "chats" not in st.session_state:
        st.session_state.chats = {
            "Chat 1": [{"role": "assistant", "content": "Hello! I am Alies AI. How can I help you today?"}]
        }
    if "current_chat" not in st.session_state:
        st.session_state.current_chat = "Chat 1"
    if "chat_counter" not in st.session_state:
        st.session_state.chat_counter = 1
    if "session_id" not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if "rag_pipeline" not in st.session_state:
        st.session_state.rag_pipeline = RAGPipeline()
    if "memory" not in st.session_state:
        st.session_state.memory = ConversationMemory(max_turns=6)
    if "llm_router" not in st.session_state:
        st.session_state.llm_router = LLMRouter()
    if "documents_loaded" not in st.session_state:
        st.session_state.documents_loaded = 0
    if "processed_files" not in st.session_state:
        st.session_state.processed_files = set()
    # Store current pipeline mode persistently
    if "pipeline_mode" not in st.session_state:
        st.session_state.pipeline_mode = "Standard Chat"
    # Image attached for the vision / OCR pipeline
    if "pending_image_bytes" not in st.session_state:
        st.session_state.pending_image_bytes = None
    if "pending_image_name" not in st.session_state:
        st.session_state.pending_image_name = None

init_state()

# ==============================================================================
# 2b. VOICE ASSISTANT (cached — loading Whisper is expensive, do it once)
# ==============================================================================
@st.cache_resource
def load_voice_assistant():
    return VoiceAssistant()

voice_assistant = load_voice_assistant()

# ==============================================================================
# 3. HELPERS
# ==============================================================================
def render_response(response: str):
    parts = re.split(r'(```[\w]*\n[\s\S]*?```)', response)
    for part in parts:
        if part.startswith('```'):
            lines = part.split('\n')
            lang = lines[0].replace('```', '').strip() or 'python'
            code = '\n'.join(lines[1:]).rstrip('`').strip()
            st.markdown(f'<div class="code-label">🖥️ {lang}</div>', unsafe_allow_html=True)
            st.code(code, language=lang)
        else:
            if part.strip():
                st.markdown(part)


def export_to_docx(messages):
    doc = DocxDocument()
    doc.add_heading(f"Chat Export — {st.session_state.current_chat}", level=1)
    for msg in messages:
        p = doc.add_paragraph()
        p.add_run(f"{msg['role'].upper()}: ").bold = True
        p.add_run(msg.get('content', ''))
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def process_uploaded_files(uploaded_files):
    """Process files and add them to RAG — returns True if any new file was loaded."""
    processor = DocumentProcessor()
    any_loaded = False
    for f in uploaded_files:
        key = f"{f.name}_{f.size}"
        if key in st.session_state.processed_files:
            continue
        suffix = f.name.split('.')[-1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{suffix}") as tmp:
            tmp.write(f.getvalue())
            tmp_path = tmp.name
        try:
            chunks = processor.process_document(tmp_path, f.name)
            ok = st.session_state.rag_pipeline.add_documents(
                chunks,
                source_type="document",
                source_name=f.name,
                session_id=st.session_state.session_id
            )
            if ok:
                st.session_state.documents_loaded += 1
                st.session_state.processed_files.add(key)
                any_loaded = True
                st.success(f"✅ Loaded: {f.name} ({len(chunks)} chunks)")
        except Exception as e:
            st.error(f"Error processing {f.name}: {e}")
        finally:
            os.unlink(tmp_path)
    return any_loaded


def process_url(url: str):
    """Scrape URL and add to RAG."""
    scraper = WebScraper()
    processor = DocumentProcessor()
    try:
        content = scraper.scrape_url(url)
        if content:
            chunks = processor.process_text(content, url)
            ok = st.session_state.rag_pipeline.add_documents(
                chunks,
                source_type="web",
                source_name=url,
                session_id=st.session_state.session_id
            )
            if ok:
                st.session_state.documents_loaded += 1
                return True, f"✅ Scraped {len(chunks)} chunks from: {url}"
            return False, "Failed to store scraped content."
        return False, "Could not extract content from URL."
    except Exception as e:
        return False, f"Scraping error: {e}"


def build_answer(user_input: str, llm_mode: str, pipeline_mode: str, image_bytes: bytes = None):
    router: LLMRouter = st.session_state.llm_router
    memory: ConversationMemory = st.session_state.memory
    session_id = st.session_state.session_id
    chat_history = memory.format_history(session_id, max_turns=4)
    sources = []

    # ── Image / OCR — attached image always goes to the vision model ─────────
    if pipeline_mode == "Image / OCR" and image_bytes is not None:
        answer = router.generate_response(
            user_input, context="", mode=llm_mode,
            image_bytes=image_bytes, chat_history=chat_history
        )

    # ── Web Search — always search live, ignore any loaded docs ──────────────
    elif pipeline_mode == "Web Search":
        searcher = WebSearchClient(max_results=5)
        results = searcher.search(user_input)
        if not results:
            answer = "I could not find any web results for your question. Please check your internet connection."
            sources = []
        else:
            context = searcher.format_results_as_context(results)
            sources = [r['url'] for r in results if r.get('url')]
            # Use dedicated web search method — avoids "context doesn't contain" response
            from src.llm_client import OllamaClient
            web_client = OllamaClient()
            answer = web_client.generate_web_search_response(user_input, context, chat_history)

    # ── RAG — only if docs are loaded ────────────────────────────────────────
    elif pipeline_mode in ("Use Uploaded Files (RAG)", "Web Scraping") \
            and st.session_state.documents_loaded > 0:
        result = st.session_state.rag_pipeline.query(
            user_input, session_id=session_id
        )
        # query() may return (answer, sources) or (answer, sources, extra)
        answer = result[0]
        sources = result[1] if len(result) > 1 else []

    # ── Standard Chat — pure LLM, no context ─────────────────────────────────
    else:
        answer = router.generate_response(
            user_input, context="", mode="smart", chat_history=chat_history
        )

    memory.add_turn(session_id, user_input, answer)
    return answer, sources


# ==============================================================================
# 4. SIDEBAR
# ==============================================================================
with st.sidebar:
    st.markdown("<h2 style='text-align:center;'>🤖 Alies AI System</h2>", unsafe_allow_html=True)

    if st.button("➕ New Chat", use_container_width=True):
        st.session_state.chat_counter += 1
        name = f"Chat {st.session_state.chat_counter}"
        st.session_state.chats[name] = [
            {"role": "assistant", "content": f"Welcome to {name}! How can I assist you?"}
        ]
        st.session_state.current_chat = name
        st.session_state.session_id = str(uuid.uuid4())
        st.session_state.processed_files = set()
        st.session_state.documents_loaded = 0
        st.session_state.pipeline_mode = "Standard Chat"
        st.session_state.pending_image_bytes = None
        st.session_state.pending_image_name = None
        st.rerun()

    st.markdown("---")
    st.markdown("<h3 style='text-align:center;'>Chat History</h3>", unsafe_allow_html=True)

    for chat_id in list(st.session_state.chats.keys()):
        col_chat, col_opts = st.columns([4, 1])
        is_active = chat_id == st.session_state.current_chat
        with col_chat:
            if st.button(
                f"💬 {chat_id}", key=f"sel_{chat_id}",
                use_container_width=True,
                type="primary" if is_active else "secondary"
            ):
                st.session_state.current_chat = chat_id
                st.rerun()
        with col_opts:
            with st.popover("⋮"):
                new_name = st.text_input("Rename", value=chat_id,
                                         key=f"rename_{chat_id}",
                                         label_visibility="collapsed")
                if st.button("✏️ Save", key=f"save_{chat_id}", use_container_width=True):
                    if new_name != chat_id and new_name not in st.session_state.chats:
                        st.session_state.chats[new_name] = st.session_state.chats.pop(chat_id)
                        if st.session_state.current_chat == chat_id:
                            st.session_state.current_chat = new_name
                        st.rerun()
                docx_data = export_to_docx(st.session_state.chats[chat_id])
                st.download_button(
                    "📄 Export DOCX", data=docx_data,
                    file_name=f"{chat_id}_export.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"exp_{chat_id}", use_container_width=True
                )
                if st.button("🗑️ Delete", key=f"del_{chat_id}",
                             use_container_width=True, type="primary"):
                    if len(st.session_state.chats) > 1:
                        del st.session_state.chats[chat_id]
                        if st.session_state.current_chat == chat_id:
                            st.session_state.current_chat = list(st.session_state.chats.keys())[0]
                        st.rerun()
                    else:
                        st.error("Cannot delete the last chat.")

    st.markdown("---")
    available_modes = st.session_state.llm_router.available_modes()
    mode_labels = [label for _, label in available_modes]
    mode_values = [val for val, _ in available_modes]
    selected_idx = st.selectbox(
        "🧠 Model / Mode",
        range(len(mode_labels)),
        format_func=lambda i: mode_labels[i]
    )
    selected_mode = mode_values[selected_idx]

    st.markdown("---")
    if st.session_state.documents_loaded > 0:
        st.success(f"📚 {st.session_state.documents_loaded} source(s) loaded")
        st.caption(f"Pipeline: **{st.session_state.pipeline_mode}**")
    else:
        st.info("📁 No documents loaded yet")

    st.markdown("---")
    if st.button("🗑️ Clear All Chats", use_container_width=True, type="primary"):
        st.session_state.chats = {
            "Chat 1": [{"role": "assistant",
                        "content": "All history cleared. Fresh session started."}]
        }
        st.session_state.current_chat = "Chat 1"
        st.session_state.chat_counter = 1
        st.session_state.session_id = str(uuid.uuid4())
        st.session_state.processed_files = set()
        st.session_state.documents_loaded = 0
        st.session_state.pipeline_mode = "Standard Chat"
        st.session_state.pending_image_bytes = None
        st.session_state.pending_image_name = None
        st.session_state.memory.clear(st.session_state.session_id)
        st.rerun()

# ==============================================================================
# 5. MAIN CHAT AREA
# ==============================================================================
header_col1, header_col2 = st.columns([3, 1])
with header_col1:
    st.subheader(f"Alies RAG Engine » {st.session_state.current_chat}")
with header_col2:
    st.markdown(
        f"<div class='backend-badge'>LLM: {selected_mode} | Pipeline: {st.session_state.pipeline_mode}</div>",
        unsafe_allow_html=True
    )

st.markdown("---")

for msg in st.session_state.chats[st.session_state.current_chat]:
    avatar = "🧑‍💻" if msg["role"] == "user" else "🤖"
    with st.chat_message(msg["role"], avatar=avatar):
        render_response(msg["content"])
        if msg.get("sources"):
            st.markdown("**Sources:**")
            for src in msg["sources"]:
                st.markdown(f'<span class="source-tag">{src}</span>',
                            unsafe_allow_html=True)

# ==============================================================================
# 6. TOOLBAR
# ==============================================================================
tool_col1, tool_col2, tool_col3, tool_col4 = st.columns([1.5, 4, 1, 1])

with tool_col1:
    with st.popover("➕ Pipeline Options", use_container_width=True):
        chosen_mode = st.radio(
            "Select Pipeline Mode:",
            ["Standard Chat", "Use Uploaded Files (RAG)", "Web Search", "Web Scraping", "Image / OCR"],
            index=["Standard Chat", "Use Uploaded Files (RAG)", "Web Search", "Web Scraping", "Image / OCR"]
                  .index(st.session_state.pipeline_mode)
        )
        # Save chosen mode to session state immediately
        if chosen_mode != st.session_state.pipeline_mode:
            st.session_state.pipeline_mode = chosen_mode
            st.rerun()

with tool_col2:
    current_pipeline = st.session_state.pipeline_mode

    if current_pipeline == "Use Uploaded Files (RAG)":
        uploaded_files = st.file_uploader(
            "Upload files",
            accept_multiple_files=True,
            type=["pdf", "docx", "pptx", "xlsx", "txt", "csv", "md"],
            label_visibility="collapsed",
            key="file_uploader_widget"
        )
        if uploaded_files:
            any_new = process_uploaded_files(uploaded_files)
            if any_new:
                st.info(f"✅ Files loaded! Now ask your question in the chat below.")

    elif current_pipeline == "Web Scraping":
        # Show already scraped sources (if any), but keep the input available
        # so the user can scrape additional URLs instead of getting stuck.
        if st.session_state.documents_loaded > 0:
            st.success(f"✅ {st.session_state.documents_loaded} source(s) loaded — ask your question below, or add another URL:")

        url_col1, url_col2 = st.columns([3, 1])
        with url_col1:
            url_input = st.text_input(
                "URL", placeholder="https://example.com",
                label_visibility="collapsed",
                key="scrape_url_input"
            )
        with url_col2:
            scrape_clicked = st.button("🌐 Scrape", use_container_width=True)

        if scrape_clicked and url_input:
            with st.spinner(f"Scraping {url_input}..."):
                success, msg = process_url(url_input)
            if success:
                st.success("✅ Done! Now type your question in the chat below.")
                st.session_state.pipeline_mode = "Web Scraping"
                st.rerun()
            else:
                st.error(msg)

    elif current_pipeline == "Web Search":
        st.info("🌐 **Web Search active** — type your question and it will search the web live.")

    elif current_pipeline == "Image / OCR":
        uploaded_image = st.file_uploader(
            "Upload an image",
            type=["png", "jpg", "jpeg", "webp"],
            label_visibility="collapsed",
            key="image_uploader_widget"
        )
        if uploaded_image is not None:
            st.session_state.pending_image_bytes = uploaded_image.getvalue()
            st.session_state.pending_image_name = uploaded_image.name
        if st.session_state.pending_image_bytes:
            st.image(st.session_state.pending_image_bytes, caption=st.session_state.pending_image_name, width=180)
            st.caption("✅ Image attached — ask your question below (e.g. \"what text is in this image?\").")
            if st.button("🗑️ Remove image", key="remove_image_btn"):
                st.session_state.pending_image_bytes = None
                st.session_state.pending_image_name = None
                st.rerun()
        else:
            st.caption("📷 Upload an image, then ask a question about it below (OCR, description, etc.).")

    else:
        st.caption("💬 Standard Chat — ask anything directly.")

with tool_col3:
    voice_audio = mic_recorder(
        start_prompt="🎙️ Audio",
        stop_prompt="⏹ Stop",
        just_once=True,
        use_container_width=True,
        key="voice_recorder"
    )

with tool_col4:
    if st.button("🧹 Clear Frame", use_container_width=True):
        st.session_state.chats[st.session_state.current_chat] = [
            {"role": "assistant", "content": "Frame cleared. Ready for new input."}
        ]
        st.session_state.memory.clear(st.session_state.session_id)
        st.rerun()

# ==============================================================================
# 7. CHAT INPUT & RESPONSE
# ==============================================================================
def process_user_message(user_input: str, spoken: bool = False):
    """Shared handler for both typed and voice-transcribed messages.
    Runs the question through the same pipeline routing (RAG / web search /
    vision / standard chat) and appends both turns to the active chat.
    Returns the assistant's answer text (used for TTS playback on voice turns).
    """
    # Auto-rename chat from first message
    current_msgs = st.session_state.chats[st.session_state.current_chat]
    if len(current_msgs) == 1 and st.session_state.current_chat.startswith("Chat "):
        new_name = (user_input[:20] + "...") if len(user_input) > 20 else user_input
        if new_name in st.session_state.chats:
            new_name += f" ({st.session_state.chat_counter})"
        st.session_state.chats[new_name] = st.session_state.chats.pop(
            st.session_state.current_chat
        )
        st.session_state.current_chat = new_name

    st.session_state.chats[st.session_state.current_chat].append(
        {"role": "user", "content": user_input}
    )
    with st.chat_message("user", avatar="🧑‍💻"):
        st.write(("🎙️ " if spoken else "") + user_input)

    with st.chat_message("assistant", avatar="🤖"):
        with st.spinner("Thinking..."):
            try:
                answer, sources = build_answer(
                    user_input,
                    selected_mode,
                    st.session_state.pipeline_mode,
                    image_bytes=st.session_state.pending_image_bytes
                )
            except Exception as e:
                answer = f"❌ Error: {e}"
                sources = []

        render_response(answer)
        if sources:
            st.markdown("**Sources:**")
            for src in sources:
                st.markdown(f'<span class="source-tag">{src}</span>',
                            unsafe_allow_html=True)

        # Voice turns get a spoken reply back
        if spoken:
            try:
                audio_output = voice_assistant.text_to_speech(answer)
                st.audio(audio_output, autoplay=True)
            except Exception as e:
                st.warning(f"Could not generate voice reply: {e}")

    st.session_state.chats[st.session_state.current_chat].append({
        "role": "assistant",
        "content": answer,
        "sources": sources
    })
    return answer


if user_input := st.chat_input("Message Alies AI..."):
    process_user_message(user_input, spoken=False)
    st.rerun()

# ── Voice input: record -> transcribe -> route through the same pipeline ──
if voice_audio:
    with st.spinner("Transcribing your voice..."):
        try:
            audio_path = voice_assistant.save_audio(
                voice_audio["bytes"], "voice_input.wav"
            )
            transcribed_text = voice_assistant.speech_to_text(audio_path)
        except Exception as e:
            transcribed_text = ""
            st.error(f"Voice transcription failed: {e}")

    if transcribed_text.strip():
        process_user_message(transcribed_text, spoken=True)
    else:
        st.warning("Didn't catch that — please try recording again.")

    voice_assistant.clean_up()
    st.rerun()